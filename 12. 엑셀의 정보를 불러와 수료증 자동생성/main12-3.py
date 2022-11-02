import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
# 파일을 읽는다
doc = docx.Document(r'12. 엑셀의 정보를 불러와 수료증 자동 생성\수료증양식.docx')
#기본폰트 속성 정하기
style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')     # 한글 폰트를 따로 설정해 준다.
style.font.size = docx.shared.Pt(12)
#문단생성하고 글씨 입력 후 폰트와 글씨크기 지정
para = doc.add_paragraph() 
run = para.add_run('\n\n')   # 문서에 단락 추가 :para.add_run()
run = para.add_run('              제 2020-0001 호\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('수  료  증') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('        성       명: 장다인\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        생 년 월 일: 2017.12.12\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20) 
run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph() 
run = para.add_run('\n\n')
run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('2021.09.19') 
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('파이썬교육기관장') 
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'12.엑셀의 정보를 불러와 수료증 자동생성\수료증결과.docx')
# docx파일생성법!

# from docx import Document

# doc = Document()
# doc.save("test.docx")
# [출처] 파이썬#93 - 사무자동화 python 으로 ms word 작성 python docx|작성자 남박사


